import pdfplumber
from docx import Document
import io
import re

# Heading keywords that signal each section (expanded for more resume formats)
SECTION_INDICATORS = {
    "skills": [
        "skills", "technical skills", "tech stack",
        "technologies", "core competencies", "key skills",
        "tools & technologies", "programming skills",
        "technical proficiency", "skill set", "skillset",
        "areas of expertise", "competencies",
    ],
    "projects": [
        "projects", "academic projects", "personal projects",
        "experience", "work experience", "project experience",
        "professional experience", "internship", "internships",
        "relevant experience", "key projects", "capstone",
    ],
    "education": [
        "education", "academic background", "qualifications",
        "academic qualifications", "academics",
        "educational background", "degrees", "certifications",
        "courses", "coursework", "relevant coursework",
    ],
    "links": [
        "github", "linkedin", "portfolio", "profiles", "social",
        "links", "contact", "websites", "online presence",
    ],
}


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # Try table extraction first for structured content
            page_text = page.extract_text(layout=True)
            if not page_text:
                page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()


def _extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # Also extract text from tables (many resumes use table layouts)
    for table in doc.tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells_text:
                parts.append(" | ".join(cells_text))
    return "\n".join(parts).strip()


def _clean_text(text: str) -> str:
    text = re.sub(r"\(cid:\d+\)", " ", text)       # PDF encoding artifacts
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", text)  # Control chars
    text = re.sub(r"-\s+", "", text)                # broken hyphenated words
    text = re.sub(r"[ \t]+", " ", text)             # collapse spaces/tabs
    text = re.sub(r"\n{3,}", "\n\n", text)          # collapse excessive newlines
    return text.strip()


def _extract_links(text: str) -> list:
    urls = re.findall(r"https?://\S+", text)
    # Catch bare linkedin / github URLs without http(s)
    for domain in ["linkedin.com", "github.com"]:
        for match in re.findall(rf"{re.escape(domain)}/\S+", text):
            full = f"https://{match}"
            if full not in urls:
                urls.append(full)
    return list(set(urls))


def _detect_sections(text: str) -> dict:
    """Split resume text into labelled sections by heading detection.

    Improved: handles UPPERCASE headings, headings with colons/dashes,
    and headings that are part of a longer line (e.g. 'SKILLS: Python, Java').
    """
    lines = text.split("\n")
    current_section = None
    section_content = {s: [] for s in SECTION_INDICATORS}

    for line in lines:
        stripped = line.strip()
        # Normalize: remove trailing colons, dashes, underscores
        normalized = re.sub(r"[\s:—\-_]+$", "", stripped).lower()
        matched = None

        for section, indicators in SECTION_INDICATORS.items():
            if any(
                normalized == kw
                or normalized.startswith(kw + " ")
                or normalized.startswith(kw + ":")
                or normalized == kw.upper()
                for kw in indicators
            ):
                matched = section
                # If the heading line also contains content after a colon,
                # capture that content too
                colon_idx = stripped.find(":")
                if colon_idx > 0 and colon_idx < len(stripped) - 1:
                    rest = stripped[colon_idx + 1:].strip()
                    if rest:
                        section_content[section].append(rest)
                break

        if matched:
            current_section = matched
        elif current_section and stripped:
            section_content[current_section].append(stripped)

    return {
        section: "\n".join(content).strip()
        for section, content in section_content.items()
        if content
    }


def detect_sections(text: str) -> dict:
    """Public wrapper: split already-extracted resume text into labelled sections.
    Used when only raw_text is available (e.g. the verification flow) to recover
    the projects/skills section text the proficiency engine needs."""
    return _detect_sections(text or "")


def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Parse a PDF or DOCX resume and return structured sections.

    Returns:
        {
            "raw_text": str,
            "skills_text": str,
            "projects_text": str,
            "education_text": str,
            "links": [str],
            "sections_detected": [str]
        }
    """
    fname = filename.lower()
    # Wrap extraction so a corrupt/unsupported file raises a clean ValueError
    # (mapped to HTTP 400 by callers) instead of leaking a pdfminer/zipfile
    # exception that surfaces as a confusing HTTP 500.
    try:
        if fname.endswith(".pdf"):
            raw = _extract_text_from_pdf(file_bytes)
        elif fname.endswith(".docx"):
            raw = _extract_text_from_docx(file_bytes)
        else:
            raise ValueError("Unsupported format. Only PDF and DOCX are allowed.")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(
            "Could not read the file — it may be corrupted, password-protected, "
            "or not a valid PDF/DOCX."
        ) from e

    raw = _clean_text(raw)
    # Reject documents with no extractable text (e.g. scanned-image PDFs) so the
    # user gets a helpful message instead of a meaningless 0 score.
    if len(raw) < 20:
        raise ValueError(
            "No readable text found in the document. If it is a scanned image, "
            "upload a text-based PDF or DOCX instead."
        )
    sections = _detect_sections(raw)
    links = _extract_links(raw)

    sections_detected = list(sections.keys())
    if links and "links" not in sections_detected:
        sections_detected.append("links")

    return {
        "raw_text": raw,
        "skills_text": sections.get("skills", ""),
        "projects_text": sections.get("projects", ""),
        "education_text": sections.get("education", ""),
        "links": links,
        "sections_detected": sections_detected,
    }
